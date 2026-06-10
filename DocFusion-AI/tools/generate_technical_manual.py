from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = REPO_ROOT / "docs"
OUTPUT_PATH = DOCS_DIR / "DocFusion-AI项目技术说明书.docx"

SCREENSHOTS = {
    "manual": REPO_ROOT / "demo" / "screenshots" / "新手指导手册.png",
    "main": REPO_ROOT / "demo" / "screenshots" / "主页面.png",
    "settings": REPO_ROOT / "demo" / "screenshots" / "提取匹配设置.png",
    "result": REPO_ROOT / "demo" / "screenshots" / "提取结果展示.png",
}

THEME = {
    "navy": "0B1527",
    "navy_soft": "14233B",
    "teal": "5EE3D2",
    "teal_soft": "E9FBF8",
    "blue": "74A8FF",
    "blue_soft": "EDF4FF",
    "text": "1F2D3D",
    "muted": "5C708A",
    "line": "D6E3F5",
    "light": "F7FAFE",
    "white": "FFFFFF",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = THEME["line"], size: int = 8) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 6, line: float = 1.4) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_run_font(run, size: float = 10.5, bold: bool = False, color: str = THEME["text"], name: str = "Aptos") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text: str, *, size: float = 10.5, bold: bool = False, color: str = THEME["text"], name: str = "Aptos"):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, name=name)
    return run


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    title = document.styles["Title"]
    title.font.name = "Aptos Display"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(THEME["navy"])

    heading1 = document.styles["Heading 1"]
    heading1.font.name = "Aptos Display"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    heading1.font.size = Pt(18)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor.from_string(THEME["navy"])

    heading2 = document.styles["Heading 2"]
    heading2.font.name = "Aptos Display"
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor.from_string("17365D")

    heading3 = document.styles["Heading 3"]
    heading3.font.name = "Aptos"
    heading3._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor.from_string("274C77")


def add_hero(document: Document) -> None:
    hero = document.add_table(rows=1, cols=1)
    hero.alignment = WD_TABLE_ALIGNMENT.CENTER
    hero.autofit = True
    cell = hero.cell(0, 0)
    set_cell_shading(cell, THEME["navy"])
    set_cell_border(cell, color=THEME["blue"], size=10)
    set_cell_margins(cell, top=260, start=220, bottom=220, end=220)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, after=10, line=1.2)
    add_text(p, "TECHNICAL MANUAL", size=9, bold=True, color=THEME["teal"], name="Aptos")

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=4, line=1.0)
    add_text(p, "DocFusion-AI", size=26, bold=True, color=THEME["white"], name="Aptos Display")

    p = cell.add_paragraph()
    set_paragraph_spacing(p, after=10)
    add_text(p, "项目概要与关键技术实现说明书", size=13.5, bold=True, color=THEME["teal"])

    p = cell.add_paragraph()
    set_paragraph_spacing(p, after=10, line=1.5)
    add_text(
        p,
        "本说明书基于当前仓库真实实现生成，重点说明系统的前端、后端、文本清洗、信息提取、信息匹配、大模型调用、"
        "关键阈值控制，以及“文本一次性过长导致匹配不稳定”的技术难题与最终解决方案。",
        size=10.5,
        color="DCE8F8",
    )

    pills = document.add_table(rows=1, cols=3)
    pills.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = [
        ("文档解析", "TXT / MD / DOCX / XLSX 统一入口"),
        ("字段抽取", "分段识别 + 逐列归并 + 溯源记录"),
        ("标准匹配", "语义映射 + 向量缓存 + 结果回退"),
    ]
    for idx, (title, desc) in enumerate(labels):
        cell = pills.cell(0, idx)
        set_cell_shading(cell, THEME["blue_soft"])
        set_cell_border(cell, color=THEME["line"])
        set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=4, line=1.2)
        add_text(p, title, size=11, bold=True, color="17365D")
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=0, line=1.25)
        add_text(p, desc, size=9.5, color=THEME["muted"])

    p = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    set_paragraph_spacing(p, before=10, after=0)
    add_text(p, "生成日期：2026-04-20", size=9.5, color=THEME["muted"])


def add_section_banner(document: Document, title: str, intro: str | None = None) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, THEME["teal_soft"])
    set_cell_border(cell, color=THEME["teal"], size=10)
    set_cell_margins(cell, top=100, start=140, bottom=100, end=140)

    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=4, line=1.1)
    add_text(p, title, size=12.5, bold=True, color="17365D")
    if intro:
        p = cell.add_paragraph()
        set_paragraph_spacing(p, after=0, line=1.35)
        add_text(p, intro, size=9.8, color=THEME["muted"])


def add_heading(document: Document, text: str, level: int = 1):
    paragraph = document.add_heading(text, level=level)
    set_paragraph_spacing(paragraph, before=10 if level == 1 else 6, after=6, line=1.2)
    return paragraph


def add_callout(document: Document, title: str, lines: Iterable[str], fill: str = THEME["light"], border: str = THEME["blue"]) -> None:
    items = list(lines)
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=border, size=10)
    set_cell_margins(cell, top=120, start=140, bottom=120, end=140)

    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=4, line=1.2)
    add_text(p, title, size=10.8, bold=True, color="17365D")

    for idx, line in enumerate(items):
        p = cell.add_paragraph()
        set_paragraph_spacing(p, after=0 if idx == len(items) - 1 else 2, line=1.35)
        add_text(p, line, size=9.8, color=THEME["text"])


def add_body_paragraph(document: Document, text: str, *, indent: bool = True) -> None:
    p = document.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line=1.45)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    add_text(p, text, size=10.5)


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        p = document.add_paragraph()
        set_paragraph_spacing(p, after=3, line=1.35)
        add_text(p, "• ", size=10.8, bold=True, color="17365D")
        add_text(p, item, size=10.3)


def add_code_block(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F7FB")
    set_cell_border(cell, color=THEME["line"], size=8)
    set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.15)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.3)
    run.font.color.rgb = RGBColor.from_string("203040")


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        set_cell_shading(cell, THEME["navy_soft"])
        set_cell_border(cell, color=THEME["blue"], size=10)
        set_cell_margins(cell, top=110, start=100, bottom=110, end=100)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=0, line=1.1)
        add_text(p, header, size=10.2, bold=True, color=THEME["white"])

    for row_idx, row_values in enumerate(rows):
        row = table.add_row()
        for col_idx, value in enumerate(row_values):
            cell = row.cells[col_idx]
            set_cell_shading(cell, THEME["white"] if row_idx % 2 == 0 else "F9FBFE")
            set_cell_border(cell, color=THEME["line"], size=8)
            set_cell_margins(cell, top=90, start=95, bottom=90, end=95)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.28)
            add_text(p, value, size=9.8, color=THEME["text"])

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)


def add_image_with_caption(document: Document, path: Path, caption: str, width: float = 6.2) -> None:
    if not path.exists():
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    set_paragraph_spacing(p, before=6, after=4, line=1.0)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=8, line=1.1)
    add_text(p, caption, size=9.5, color=THEME["muted"])


def add_page_break(document: Document) -> None:
    p = document.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_contents(document: Document) -> None:
    add_section_banner(document, "目录", "本说明书按“从产品目标到技术难点”的顺序展开，便于答辩、汇报和交付。")
    toc_rows = [
        ["1 项目概述", "系统定位、用户使用链路、项目价值"],
        ["2 系统架构与项目目录", "前后端关系、目录分层、数据流转"],
        ["3 前端技术实现", "上传、轮询、预览、结果展示、导出、溯源交互"],
        ["4 后端技术实现", "FastAPI 路由、后台流水线、缓存、恢复、数据库"],
        ["5 文档解析与文本清洗", "多格式解析、标准化、表格视图构建、清洗规则"],
        ["6 信息提取技术实现", "字段语义初始化、分批识别、逐列归并、结果落库"],
        ["7 信息匹配技术实现", "标准字段映射、向量匹配、跳过机制、溯源回写"],
        ["8 大模型调用情况", "调用入口、提示词设计、JSON 约束、回退策略"],
        ["9 技术难题与最终解决", "长文本失配、大表格膨胀、归并歧义、恢复机制"],
        ["10 结论与后续优化方向", "当前能力总结、工程可扩展方向"],
    ]
    add_table(document, ["章节", "说明"], toc_rows, widths=[2.2, 4.2])
    add_callout(
        document,
        "阅读建议",
        [
            "如果用于项目答辩，建议重点阅读第 1、6、8、9 章，能够完整解释“为什么系统可行、为什么抽取得到结果、为什么遇到难题后还能落地”。",
            "如果用于后续研发交接，建议重点阅读第 3、4、5、6、7 章，里面已经把关键代码路径、阈值、回退逻辑和数据结构讲清楚。",
        ],
        fill="F7FBFF",
        border=THEME["blue"],
    )


def add_project_overview(document: Document) -> None:
    document.add_heading("1 项目概述", level=1)
    add_body_paragraph(
        document,
        "DocFusion-AI 是一个面向异构文档的结构化信息抽取与融合系统。系统的核心思想不是先设计固定模板、再让用户去适配模板，"
        "而是让用户先定义自己需要的结果列，再由系统围绕这些列执行解析、抽取、归并、标准化匹配与结果溯源。"
        "这种“字段先行”的产品设计，直接降低了业务侧接入门槛，也决定了后端必须具备较强的动态字段理解能力。"
    )
    add_body_paragraph(
        document,
        "当前主线版本支持 TXT、MD、DOCX、XLSX 四类文件。用户在前端输入列名后，可一次性上传多个文件；"
        "系统会为每个文件单独创建任务，在后台完成解析与抽取，再将结果以按文件分组的结果卡片形式展示。"
        "每条结果不仅有字段值，还有段落编号、表格坐标、来源文本等溯源信息，因此项目并非只关注“抽得出来”，也同时关注“能不能解释结果为什么这样来”。"
    )
    add_callout(
        document,
        "项目的核心产品目标",
        [
            "把非结构化文本、半结构化表格和多格式办公文档统一纳入一条处理流水线。",
            "让用户用业务语言定义输出列，而不是要求用户先理解底层模板结构。",
            "保证结果可回看、可追溯、可导出，不把系统做成只能“看答案”的黑盒。",
        ],
    )
    add_heading(document, "1.1 用户使用链路", level=2)
    add_table(
        document,
        ["步骤", "前端动作", "后端动作", "产出"],
        [
            ["步骤 1", "在页面左侧配置输出列名与槽位", "接收 extract_config 并做规范化", "生成字段任务定义"],
            ["步骤 2", "上传一个或多个文件", "保存文件、计算哈希、创建 Task、命中缓存则直接复用", "任务编号 task_id"],
            ["步骤 3", "前端轮询进度", "后台按 parse -> extract -> match 顺序执行", "阶段进度、错误信息"],
            ["步骤 4", "查看抽取结果、表格、原文", "返回 fields 预览数据与 trace 信息", "结构化结果卡片"],
            ["步骤 5", "点击字段或单元格查看来源", "查询字段/记录/单元格来源", "段落、表格、证据文本"],
            ["步骤 6", "导出结果表", "前端按当前字段顺序导出 xlsx", "可交付结果文件"],
        ],
        widths=[1.0, 2.0, 2.3, 1.5],
    )
    add_heading(document, "1.2 项目价值", level=2)
    add_bullets(
        document,
        [
            "对业务侧而言，系统把“理解文档、整理字段、回填结果表”三件事整合成一个界面动作，显著降低人工整理工作量。",
            "对研发侧而言，系统通过任务状态、字段溯源、向量缓存和结果回退，保证整个流水线具备工程稳定性，而不是只在理想测试样例下工作。",
            "对汇报与答辩而言，系统可以明确说明：它如何解析文件、如何调用模型、如何控制长文本风险、如何解决多记录归并问题，以及为什么最终结果可追踪。"
        ],
    )


def add_architecture(document: Document) -> None:
    document.add_heading("2 系统架构与项目目录", level=1)
    add_body_paragraph(
        document,
        "系统整体采用“静态前端 + FastAPI 后端 + SQLite 任务存储 + 规则与模型混合引擎”的轻量架构。"
        "根目录通过 start.py 统一启动，FastAPI 在 backend/app/main.py 中注册 API 路由，并直接挂载 frontend 目录作为静态站点。"
        "因此浏览器与后端服务默认同源，前端请求路径可以直接使用 window.location.origin 作为基础地址，部署与调试成本都较低。"
    )
    add_code_block(
        document,
        "Browser UI\n"
        "  -> /upload/batch\n"
        "  -> Task + progress\n"
        "  -> parse (文档结构解析)\n"
        "  -> extract (字段识别与逐列归并)\n"
        "  -> match (标准字段映射，可跳过)\n"
        "  -> /fields/{task_id} 展示结果与溯源\n"
        "  -> 浏览器导出 xlsx"
    )
    add_heading(document, "2.1 目录分层说明", level=2)
    add_code_block(
        document,
        "backend/                 FastAPI 后端与业务实现\n"
        "backend/app/api/         路由层与实际生效的抽取/匹配引擎\n"
        "backend/app/services/    通用服务，如 llm_client、document_parser、资产文件\n"
        "backend/app/db/          数据库模型、初始化与迁移兼容\n"
        "backend/storage/         上传文件、日志、SQLite 数据库等运行时产物\n"
        "frontend/                静态前端页面、样式和浏览器侧逻辑\n"
        "docs/                    文档与开发记录\n"
        "demo/screenshots/        演示截图\n"
        "test_data/               测试样例与历史样本\n"
        "legacy/                  旧版本或历史实验代码\n"
        "start.py                 启动入口"
    )
    add_heading(document, "2.2 关键模块职责", level=2)
    add_table(
        document,
        ["模块", "职责", "实现要点"],
        [
            ["frontend/index.html + MainPage.js", "浏览器交互层", "字段配置、批量上传、轮询进度、结果展示、预览、导出、溯源弹窗"],
            ["backend/app/main.py", "应用入口", "注册健康检查、上传、任务、解析、抽取、字段、匹配等路由，并托管静态前端"],
            ["backend/app/api/upload.py", "任务创建层", "保存文件、计算哈希、缓存命中、线程池提交后台流水线"],
            ["backend/app/api/parse.py", "解析层", "调用 DocumentParser 读取原文、段落、表格与 table_views"],
            ["backend/app/api/_extract_engine.py", "核心抽取层", "字段语义初始化、分批识别、候选验证、逐列归并、决策回退"],
            ["backend/app/api/match.py", "标准匹配层", "抽取键值对、判定是否适合匹配、语义映射与溯源构建"],
            ["backend/app/api/fields.py", "结果包装层", "预览裁剪、表格 trace 命中、字段/记录来源查询"],
            ["backend/app/db/models.py", "持久化层", "Task、DocumentField、ExtractedEntity、FieldEmbeddingCache"],
        ],
        widths=[1.6, 1.5, 3.1],
    )
    add_heading(document, "2.3 数据流转机制", level=2)
    add_bullets(
        document,
        [
            "原始文件流：前端上传文件 -> 后端保存到 backend/storage/uploads -> 在 Task 中记录 file_path、file_hash、processor_version、extract_config。",
            "结构化中间态：解析结果写入 Task.result；抽取结果写入 Task.extract_result；匹配结果写入 Task.match_result，并在 Task.result 中做合并保留。",
            "结果查询态：/fields/{task_id} 对大结果做裁剪，只返回预览；完整结果仍保存在 Task 与 ExtractedEntity 中，便于追踪与后续扩展。",
            "兼容与迁移：paths.py 会把 legacy 路径下的上传目录、日志目录、数据库位置迁移到统一 storage 布局，降低版本迭代时的路径混乱风险。",
        ],
    )


def add_frontend(document: Document) -> None:
    document.add_heading("3 前端技术实现", level=1)
    add_body_paragraph(
        document,
        "前端采用静态页面方案，而不是引入复杂的前端框架。主入口 frontend/index.html 通过 ES Module 加载 MainPage.js，"
        "页面交互和状态管理全部在浏览器端原生完成。这样做的优势是：部署简单、与 FastAPI 集成直接、构建链条轻、答辩演示稳定。"
        "配合 Mammoth 和 SheetJS，前端还能在本地浏览器中直接预览 Word 和 Excel。"
    )
    add_heading(document, "3.1 页面状态与交互组织", level=2)
    add_table(
        document,
        ["状态字段", "作用", "技术细节"],
        [
            ["files", "待上传文件队列", "去重校验文件名、大小、修改时间；限制扩展名为 .txt/.md/.docx/.xlsx"],
            ["extractFields", "结果列配置", "每列包含 id、slot、label、type、enabled，用于构造 extract_config"],
            ["results", "任务结果列表", "保存 taskId、pending、cached、task 快照、data 结果和错误信息"],
            ["preview", "预览弹窗", "按文件类型调用不同预览器"],
            ["trace", "溯源弹窗", "展示字段、记录、单元格对应来源"],
            ["status", "顶部状态提示", "统一反馈上传、处理中、成功、失败、警告等状态"],
        ],
        widths=[1.3, 1.3, 3.6],
    )
    add_body_paragraph(
        document,
        "字段配置区不是静态写死的表单，而是一个动态列构建器。用户新增列时需要给出列名与槽位，槽位取值包括 category、indicator、value、unit、time、yoy。"
        "槽位不直接决定最终结果名称，但会影响后端的字段语义推断、数值检查、时间检查和后续归并策略。"
    )
    add_heading(document, "3.2 上传与进度轮询", level=2)
    add_body_paragraph(
        document,
        "点击上传后，前端先将当前列配置整理成 extract_config：仅保留启用列、按界面顺序输出 field_name、slot、type、visible。"
        "随后调用 /upload/batch。上传成功后并不立即去取结果，而是把每个 task_id 放入 results 列表，启动 monitorTaskEntries 按任务轮询 /tasks/{task_id}/progress。"
        "这套设计把“任务状态”和“结果载荷”分开，避免处理尚未完成时直接拉大体积结果。"
    )
    add_bullets(
        document,
        [
            "MAX_POLL_ERRORS 为 20，连续失败达到阈值后会终止轮询并向界面报告错误。",
            "当任务状态进入 matched 或 extracted 且 extract_status=success 时，前端才调用 /fields/{task_id} 获取可展示结果。",
            "若页面刷新，localStorage 会保存最近最多 12 个任务的 taskId 与简要状态，前端重开时会执行恢复逻辑，继续轮询或直接取结果。",
        ],
    )
    add_heading(document, "3.3 预览、导出与溯源体验", level=2)
    add_table(
        document,
        ["能力", "实现方式", "说明"],
        [
            ["TXT/MD 预览", "File.text()", "直接读取文本并做 HTML 转义"],
            ["DOCX 预览", "mammoth.convertToHtml", "把 docx 转换为浏览器可显示的 HTML"],
            ["XLSX 预览", "SheetJS xlsx.full.min.js", "读取工作表并输出 HTML 表格"],
            ["单任务导出", "XLSX.utils.aoa_to_sheet + writeFile", "按当前字段顺序导出结果表"],
            ["全部结果导出", "多 sheet 工作簿", "每个成功任务一个 sheet，文件名统一按日期生成"],
            ["字段溯源", "getFieldSource / getRecordSource", "点击字段值打开弹窗，显示来源段落、表格坐标和证据文本"],
        ],
        widths=[1.2, 2.0, 3.0],
    )
    add_body_paragraph(
        document,
        "前端的溯源交互不仅支持字段级，还支持记录级与单元格级。对于解析出来的 table_views，前端会给每个单元格绑定 trace-cell 事件；"
        "若后端已经将某个抽取结果回连到该单元格，则界面会高亮显示命中并弹出关联记录信息。"
        "这让系统从单纯的“表格输出器”变成了“可解释的抽取工作台”。"
    )
    add_heading(document, "3.4 前端为何不使用复杂框架", level=2)
    add_bullets(
        document,
        [
            "项目当前更重视文档处理和算法闭环，而不是复杂的前端组件复用，因此静态页面方案可以把工程精力集中到核心问题。",
            "FastAPI 直接托管静态目录，避免额外的打包、代理、跨域和前端部署链路。",
            "所有核心交互已经具备：配置、上传、状态、结果、导出、溯源。对比赛型或研究型项目而言，这种实现更稳、更便于快速迭代。",
        ],
    )
    add_image_with_caption(document, SCREENSHOTS["settings"], "图 1 提取列配置与上传工作区")
    add_image_with_caption(document, SCREENSHOTS["result"], "图 2 结果展示、标准匹配与表格溯源界面")


def add_backend(document: Document) -> None:
    document.add_heading("4 后端技术实现", level=1)
    add_body_paragraph(
        document,
        "后端基于 FastAPI 实现，应用入口位于 backend/app/main.py。系统会注册 health、upload、tasks、parse、extract、fields、match、trace 等路由，"
        "同时挂载前端静态目录。CORS 被设置为 allow_origins=['*']，这既保证了同源静态部署，也兼容后续分离式调试。"
    )
    add_heading(document, "4.1 上传受理与缓存复用", level=2)
    add_body_paragraph(
        document,
        "upload.py 是整个流水线的起点。文件上传后，后端会先把二进制内容写入 storage/uploads，然后计算 SHA-256 哈希。"
        "同一份文件如果在相同的 processor_version（当前代码为 2026.04.16.v5）和相同 extract_config 下再次上传，"
        "系统会直接命中缓存结果而不是重新跑完整流水线。这一点对批量调试和答辩演示非常重要，因为它显著减少了重复调用模型与重复解析的成本。"
    )
    add_table(
        document,
        ["后端机制", "实现位置", "说明"],
        [
            ["文件存储", "upload.py::save_upload_content", "按原文件名保存，若同名则拼接哈希前缀避免覆盖"],
            ["缓存命中", "upload.py::find_cached_task", "按 file_hash + processor_version + extract_config 复用结果"],
            ["线程池执行", "PIPELINE_EXECUTOR", "默认 4 个 worker，后台执行 parse/extract/match"],
            ["任务恢复", "recover_task_pipeline_if_needed", "任务长时间停滞后自动重新入队"],
            ["失败兜底", "mark_task_failed", "统一设置 status、错误信息和 100% 失败进度"],
        ],
        widths=[1.3, 2.0, 2.9],
    )
    add_heading(document, "4.2 后台流水线与进度状态", level=2)
    add_body_paragraph(
        document,
        "真正的处理顺序是 parse -> extract -> match。上传接口只负责创建 Task 并把任务提交到线程池，"
        "实际执行发生在 run_task_pipeline 中。任务表不仅有总状态 status，还有 parse_status、extract_status、match_status 三个阶段状态，"
        "以及 progress_stage、progress_percent、progress_message 等细粒度字段。前端因此可以明确展示“正在解析文件结构”“正在按行识别并逐列归并字段”“正在执行标准字段匹配”等真实进度，而不是只有一个模糊的 loading。"
    )
    add_bullets(
        document,
        [
            "解析阶段完成后，进度大约推进到 25%。",
            "抽取阶段开始前进度设置到 28%，行级批次识别与列归并在 38%~92% 区间推进。",
            "匹配阶段从 94% 进入，完成后到 100%；若不适合匹配则会标记 skipped 而不是伪装为成功。",
        ],
    )
    add_heading(document, "4.3 数据库设计与兼容迁移", level=2)
    add_table(
        document,
        ["数据表", "核心字段", "作用"],
        [
            ["tasks", "file_hash / extract_config / result / extract_result / match_result / progress_*", "保存任务全流程状态与完整结果 JSON"],
            ["document_fields", "category / indicator / value / unit / time / yoy / source_*", "兼容旧接口，同时保存主结果和来源定位"],
            ["extracted_entities", "record_id / field_name / field_value / source_* / confidence", "把抽取结果按字段拆开保存，便于模板填表和溯源"],
            ["field_embedding_cache", "model_name / field_hash / embedding", "缓存字段名向量，降低重复语义匹配成本"],
        ],
        widths=[1.4, 2.7, 2.1],
    )
    add_body_paragraph(
        document,
        "数据库采用 SQLite，默认路径为 backend/storage/app.db。database.py 在 init_db 中不仅会 create_all，还会对旧表执行列检查与 ALTER TABLE，"
        "确保历史数据结构不至于因为字段增加而直接失效。此外，paths.py 会自动迁移旧版 uploads/logs/app.db 到新的 storage 目录布局，"
        "体现出本项目在重构过程中对兼容性的重视。"
    )
    add_heading(document, "4.4 后端的工程取舍", level=2)
    add_bullets(
        document,
        [
            "结果以 JSON 字符串形式原样写回 Task，而不是拆成过于细碎的多表关系，目的是先保证流水线完整与结果可回放。",
            "同时用 ExtractedEntity 保存拆分后的字段级记录，兼顾后续填表、追踪、匹配和统计。",
            "线程池模型足以覆盖当前文档处理型任务，避免引入更复杂的消息队列与 worker 编排成本。",
        ],
    )


def add_parser_and_cleaning(document: Document) -> None:
    document.add_heading("5 文档解析与文本清洗", level=1)
    add_body_paragraph(
        document,
        "文本清洗与结构解析是整个系统最基础、也最关键的一层。因为后续所有模型判断、正则兜底、主键归并和标准匹配都建立在“段落、表格、单元格已经被干净抽出”的前提之上。"
        "如果这一层处理粗糙，后面的模型再强也会因为输入质量不稳定而产生大量噪声。"
    )
    add_heading(document, "5.1 多格式解析策略", level=2)
    add_table(
        document,
        ["文件类型", "解析方式", "关键实现"],
        [
            ["DOCX", "python-docx 优先，异常时回退手工 XML 解析", "既读取段落也读取表格；必要时解压 zip 并解析 word/document.xml"],
            ["XLSX", "openpyxl 优先，失败时回退 python-calamine", "统一输出表格行、表头、单元格定位与预览段落"],
            ["MD", "按文本处理", "保留段落结构，供后续批次识别"],
            ["TXT", "按文本处理", "原样读取后规范化清洗"],
        ],
        widths=[1.0, 2.3, 2.9],
    )
    add_body_paragraph(
        document,
        "DocumentParser.parse 的标准输出是统一的：doc_id、doc_type、paragraphs、paragraph_items、tables、table_views、raw_text。"
        "也就是说，不同格式的文件最终都会被转换成同一种中间态。后续抽取引擎无需再关心文件原始类型，只需消费统一的数据结构。"
    )
    add_heading(document, "5.2 文本清洗规则", level=2)
    add_bullets(
        document,
        [
            "去除控制字符：删除不可见控制符，避免模型输入出现异常字符噪声。",
            "压缩空白：把连续空白、换行、制表符折叠为单个空格，提高比较与匹配稳定性。",
            "统一标点：将全角冒号统一为半角冒号，便于后续键值对识别与正则提取。",
            "去除脚注标号：例如 [1] 这类脚注引用在统计公报中大量存在，会干扰数值抽取，因此在抽取引擎中会先删掉。",
            "空行与空单元格过滤：仅保留含有效内容的段落和行，减少无意义输入。"
        ],
    )
    add_heading(document, "5.3 表格视图与单元格定位", level=2)
    add_body_paragraph(
        document,
        "对 Excel 和 Word 表格，系统不仅保存二维数组形式的 tables，还会额外构造 table_views。"
        "在 table_views 中，每个单元格都带有 row_index、col_index、locator（如 A1、B3）和 value，"
        "并通过首个“至少有两个非空值”的行推断表头。这样做有两个直接收益：第一，前端可以把表格直接渲染出来；第二，后端和前端都能把抽取结果回连到具体单元格。"
    )
    add_heading(document, "5.4 为什么要把表格行转成段落", level=2)
    add_body_paragraph(
        document,
        "抽取引擎并不只依赖原始 paragraph_items。对于 table_views，系统会把每一行改写成“表头: 值 | 表头: 值”的段落化文本，"
        "然后再放进后续批次识别。这样可以让模型用统一的段落抽取框架同时处理普通文本和表格行，而不必为每种输入形态分别维护完全不同的提示词和后处理逻辑。"
    )
    add_heading(document, "5.5 大文件控制策略", level=2)
    add_callout(
        document,
        "避免一次性把大表格全部送入模型",
        [
            "代码中通过 MAX_XLSX_ROW_PARAGRAPHS=120 限制 Excel 转换出的段落数。",
            "若表格行数过大，算法会保留前 40 行，再对剩余行做均匀采样，而不是简单截断。",
            "这一策略保证前部关键结构不丢失，同时避免数千行表格直接膨胀成无法稳定处理的长文本。",
        ],
        fill="FFF8E8",
        border="E0B64C",
    )


def add_extraction(document: Document) -> None:
    document.add_heading("6 信息提取技术实现", level=1)
    add_body_paragraph(
        document,
        "信息提取是 DocFusion-AI 的核心。当前实际运行的代码位于 backend/app/api/_extract_engine.py。"
        "这部分并不是“把全文喂给模型，让模型一次性输出整张结果表”，而是一个多阶段、可回退、带置信度与溯源的混合流程。"
    )
    add_heading(document, "6.1 从前端列名构建字段任务", level=2)
    add_body_paragraph(
        document,
        "前端传入的 extract_config 先经过 _extract_config.py 规范化，推断 slot 和 field type。"
        "随后 extract() 会调用 build_field_tasks_from_frontend：它把用户列名、槽位、现有描述和文档摘要（截断到 14000 字符）发给大模型，"
        "让模型只做一件事：为每个列名补齐内部 key、description 和 aliases。"
        "这个阶段的目标不是直接抽值，而是把用户的业务列名转换成后续可执行的“字段任务定义”。"
    )
    add_table(
        document,
        ["字段任务属性", "含义", "用途"],
        [
            ["field_name", "用户在前端看到的列名", "最终结果输出列"],
            ["name", "内部 snake_case 键名", "用于模型返回和内部索引"],
            ["slot", "字段角色", "影响数值/时间/分类等校验规则"],
            ["aliases", "同义表达", "帮助命中原文中的不同写法"],
            ["description", "该列真正想抽取什么", "作为模型理解字段意图的补充语义"],
            ["visible", "是否在前端展示", "控制结果列可见性"],
        ],
        widths=[1.4, 2.0, 2.6],
    )
    add_heading(document, "6.2 分批识别行候选，而不是一次性整表抽取", level=2)
    add_body_paragraph(
        document,
        "extract_row_candidates 会先把统一段落列表切成批次。当前默认参数是每批最多 10 个段落或 7000 个字符。"
        "每一批都单独构造 JSON 提示词，要求模型只在同一 paragraph_id 内判断字段，不允许跨段借信息、不允许补全、不允许猜测。"
        "若模型不可用或返回空结果，则立刻回退到 fallback_extract_row_candidates，用正则和槽位规则做最小可用提取。"
    )
    add_bullets(
        document,
        [
            "时间类字段用 DATE_PATTERN 检查；数值类字段用 NUMBER_PATTERN 抽取 number 和 unit。",
            "validate_row_field 会验证 value 是否真正在段落或 evidence 中出现，防止模型返回“整行复制”的伪字段。",
            "MIN_ROW_FIELD_CONFIDENCE=0.72，用于过滤低置信度字段，避免噪声被后续归并放大。",
            "对于表格行来源，source_kind、source_table_id、source_row、source_locator 等定位信息会一并保留。",
        ],
    )
    add_heading(document, "6.3 候选验证与字段锚点校验", level=2)
    add_body_paragraph(
        document,
        "仅仅从模型输出拿到 value 还不够，系统还要确认这个 value 不是误提、错提或整段复制。"
        "因此 validate_row_field 会做多重校验：一是 value 不能为空也不能是“略、未提及”之类的占位词；"
        "二是 value 必须能在原段落或 evidence 中找到；三是字段名或别名要能在段落语境中找到锚点；"
        "四是时间列必须长得像日期，数值列必须至少包含数字。经过这一层筛选后，留下的候选字段才会进入归并阶段。"
    )
    add_heading(document, "6.4 逐列归并：先建行，再补列", level=2)
    add_body_paragraph(
        document,
        "系统并不是先识别完整行，而是先用第一列作为种子行。build_first_column_rows 会基于首列候选创建初始记录，"
        "然后对后续每一列执行 collect_row_field_options -> decide_field_option_for_row -> apply_selected_option_to_row。"
        "这种策略的实质，是把复杂的“整行多字段组合优化”拆成“当前行针对当前列的候选选择问题”，显著降低了模型决策难度。"
    )
    add_table(
        document,
        ["机制", "关键阈值/逻辑", "工程意义"],
        [
            ["候选评分", "MIN_FIELD_OPTION_SCORE=0.9", "先用启发式筛掉明显不靠谱的选项"],
            ["直接自动采用", "AUTO_SELECT_ROW_FIELD_SCORE=7.2", "候选与主键/来源强一致时不再浪费模型调用"],
            ["LLM 决策采用", "MIN_ROW_FIELD_DECISION_CONFIDENCE=0.64", "仅在模型足够有把握时使用其选择"],
            ["主键归并 LLM 采用", "MIN_ROW_MATCH_CONFIDENCE=0.82", "避免把不同记录误并到同一行"],
            ["可解释记录", "__field_options__ / __decision_trace__", "保留备选项、原因与置信度，便于调试和展示"],
        ],
        widths=[1.4, 1.9, 2.8],
    )
    add_heading(document, "6.5 混合策略：启发式 + 语义提示 + 模型裁决", level=2)
    add_body_paragraph(
        document,
        "每个字段候选都会先计算启发式得分，得分依据包括：是否与当前行主键值一致、是否处于同一表格行、是否在同一段落、"
        "上下文窗口内是否出现当前行已知键值、字段语义提示是否把该短语判为目标字段。"
        "其中语义提示由 FieldAssistMatcher 提供，它可选加载 shibing624/text2vec-base-chinese 计算语义相似度；"
        "如果向量模型不可用，就回退到字符重叠 + SequenceMatcher。只有当候选之间仍然难分高下时，才调用 LLM 做最终裁决。"
    )
    add_heading(document, "6.6 结果落库与后续接口服务", level=2)
    add_body_paragraph(
        document,
        "抽取完成后，finalize_rows 会把内部 row 结构转换为前端友好的结果：每条记录包含 record_id、各输出列字段值、"
        "__sources__、__field_options__、__decision_trace__ 和 __key_fields__。随后 extract.py 会同步写入 Task.extract_result，"
        "并调用 save_document_field 与 save_extracted_entities 将核心字段和字段级实体落库。"
        "这一步让结果既能直接展示，又能在后续做模板填表、字段追踪和二次统计。"
    )


def add_matching(document: Document) -> None:
    document.add_heading("7 信息匹配技术实现", level=1)
    add_body_paragraph(
        document,
        "DocFusion-AI 的“信息匹配”不是抽取主链路的唯一出口，而是一个标准字段映射层。"
        "抽取层已经能生成用户定义的结果表；匹配层的职责是把一些更偏业务键值对文档中的字段，进一步映射到标准字段字典中。"
        "因此 match 路由具备“能做则做，不适合就跳过”的设计，而不是强制所有文件都做标准化。"
    )
    add_heading(document, "7.1 何时进入匹配", level=2)
    add_body_paragraph(
        document,
        "match.py 会先从 parse_data 中抽取输入项：正文里的“字段: 值”对，以及表格里“表头 -> 单元格值”的组合。"
        "然后根据文档形态判断是否适合匹配。若是明显的大型表格型数据（例如 xlsx 且行数 >= 12 或列数 >= 10），"
        "系统会直接返回 skipped，理由是这类文件更适合保留 extract + table trace，而不是强行压到少量标准字段上。"
    )
    add_heading(document, "7.2 标准字段匹配器", level=2)
    add_body_paragraph(
        document,
        "标准匹配器 FieldSemanticMatcher 读取 services/assets/field_mapping.json。"
        "该字典目前包含项目名称、负责人、单位名称、姓名、电话、专业、金额、部门、人数等标准字段及其中文同义词。"
        "匹配流程先做规范化文本比较，若 extracted_key 的规范化形式在 reverse_dict 中能直接命中，则直接返回 1.0 分；"
        "否则进入语义相似度计算。"
    )
    add_table(
        document,
        ["匹配层能力", "实现方式", "效果"],
        [
            ["直接同义命中", "normalize_key + reverse_dict", "同义词字典稳定、速度快"],
            ["向量匹配", "sentence-transformers + cosine_similarity", "适合表达变化较大的字段名"],
            ["回退匹配", "SequenceMatcher + 字符集合重叠", "模型不可用时仍可运行"],
            ["规则校验", "_rule_check", "如 amount/quota 必须带数字，避免语义相近但类型错误"],
            ["向量缓存", "FieldEmbeddingCache", "避免重复对同一字段名编码"],
        ],
        widths=[1.4, 2.0, 2.7],
    )
    add_heading(document, "7.3 匹配结果为什么还能溯源", level=2)
    add_body_paragraph(
        document,
        "即使完成了标准字段映射，系统也不会丢掉来源信息。build_matched_trace_map 会把 matched_result 中的每个标准字段值回查到原始 input_items，"
        "保留 source_key、source_kind、source_paragraph、source_table_id、source_locator 等信息。"
        "前端点击匹配结果时，可以继续看到“这个标准字段是从哪个原始字段、哪段文本或哪张表格里来的”。"
    )
    add_heading(document, "7.4 匹配层的工程定位", level=2)
    add_bullets(
        document,
        [
            "对于表格型统计数据，主要价值来自 extract 结果和单元格 trace，match 更多是辅助层。",
            "对于业务申请表、报名表、项目申报类文档，match 能把“联系电话/手机/咨询电话”等不同表达统一到标准字段，提高结果复用性。",
            "由于 match 支持 skipped 和 failed 回退，系统不会因为标准字段字典不全而破坏主流程结果。"
        ],
    )


def add_llm_usage(document: Document) -> None:
    document.add_heading("8 大模型调用情况", level=1)
    add_body_paragraph(
        document,
        "项目的大模型调用统一经过 services/llm_client.py。该模块使用 OpenAI 兼容 SDK，默认 base_url 指向 DashScope 兼容接口，默认模型为 qwen-plus。"
        "系统优先读取 EXTRACT_API_KEY，其次读取 DASHSCOPE_API_KEY，再次读取 OPENAI_API_KEY。只要满足“有 key 且安装了 openai 包”，就会构建客户端；否则所有模型相关步骤自动进入回退模式。"
    )
    add_heading(document, "8.1 大模型实际调用的四个场景", level=2)
    add_table(
        document,
        ["调用场景", "输入", "输出", "失败后的回退"],
        [
            ["字段语义初始化", "前端列名 + 截断后的文档摘要", "内部 key、字段描述、别名", "直接使用前端列名与规则推断 slot"],
            ["批次行候选识别", "最多 10 个段落 / 7000 字符", "每段的字段 records JSON", "使用正则与槽位规则做 fallback_extract_row_candidates"],
            ["列级候选决策", "当前行已知主键 + 候选项列表", "selected_option_id / alternatives / confidence", "使用启发式最高分回退或不填充"],
            ["主键归并歧义判断", "候选字段值 + 多个候选行主键", "row_id / confidence", "采用最高分启发式候选行或返回 NONE"],
        ],
        widths=[1.7, 1.9, 1.9, 1.9],
    )
    add_heading(document, "8.2 提示词设计原则", level=2)
    add_bullets(
        document,
        [
            "每个提示词都强制要求只输出 JSON，不允许解释性废话，降低解析难度。",
            "字段识别提示词强调“每个 paragraph_id 独立判断，不能借用其他段落信息”，防止跨段串值。",
            "列级决策提示词强调“match_hint 只作辅助，最终必须依据字段语义和当前行主键判断”，防止模型被高相似词误导。",
            "主键归并提示词强调“只能从候选行中选一个或 NONE”，防止模型随意虚构新记录。",
            "temperature 基本固定为 0.0，llm_client 默认 0.1；整体目标是稳定性优先而不是创造性。",
        ],
    )
    add_heading(document, "8.3 为什么模型输出要做二次校验", level=2)
    add_body_paragraph(
        document,
        "模型输出在本系统中永远不是终局答案，而是候选答案。"
        "一方面，llm_client 会先尝试提取 markdown fenced JSON，再 safe_load_json，解析失败直接返回空字典；"
        "另一方面，抽取引擎和匹配引擎还会继续做 value 是否存在于原文、字段锚点是否命中、置信度是否过低、是否为整行转储等校验。"
        "这说明项目没有把 LLM 当成全知全能黑盒，而是把它嵌入到一个可验证的工程系统里。"
    )
    add_heading(document, "8.4 关键环境变量", level=2)
    add_table(
        document,
        ["环境变量", "默认值/含义", "说明"],
        [
            ["EXTRACT_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1", "大模型服务地址"],
            ["EXTRACT_MODEL", "qwen-plus", "主要对话模型"],
            ["EXTRACT_API_KEY / DASHSCOPE_API_KEY / OPENAI_API_KEY", "按优先级选择", "任一存在即可初始化客户端"],
            ["EXTRACT_FIELD_MATCH_MODEL", "shibing624/text2vec-base-chinese", "字段语义提示模型"],
            ["EXTRACT_ENABLE_FIELD_MATCH_EMBEDDING", "0", "开启后尝试加载字段匹配向量模型"],
        ],
        widths=[2.4, 2.4, 1.9],
    )


def add_challenges(document: Document) -> None:
    document.add_heading("9 技术难题与最终解决", level=1)
    add_heading(document, "9.1 难题一：一次性传入太多文本，无法稳定匹配", level=2)
    add_body_paragraph(
        document,
        "这是项目实现过程中最核心的技术难题。最初如果把整篇文档、整张表甚至多个表一起交给模型，"
        "会出现三类问题：第一，输入过长导致模型上下文压力增大，返回内容经常混入多条记录；"
        "第二，模型倾向于复制整段原文或整行表格，而不是给出最小可用字段值；"
        "第三，后续主键归并几乎无法进行，因为一条模型输出里已经混入了多个对象、多个地区或多个时间点。"
    )
    add_callout(
        document,
        "最终解决方案不是“换一个更大的模型”，而是“改变输入组织方式”",
        [
            "字段语义初始化阶段只截取前 14000 个字符，不让模型在准备阶段就吃掉全部文档。",
            "行候选识别阶段采用 10 段落 / 7000 字符的批次切分，每个 batch 独立判断。",
            "Excel 等大表只保留前 40 行并对剩余行均匀采样，总段落数控制在 120 以内。",
            "抽取与归并分离：先识别字段候选，再逐列归并，不再让模型一次性完成所有工作。",
            "对每个字段值都要求 evidence 来自同段，且要通过规则校验，防止长文本错配被放大。",
        ],
        fill="FFF4E9",
        border="E49C4C",
    )
    add_body_paragraph(
        document,
        "这一套调整后，系统从“整文档大一统抽取”转变成“结构解析 -> 分批识别 -> 逐列决策 -> 主键归并”的分治流水线。"
        "这正是项目最终能够稳定工作的关键，也是答辩时最值得重点讲清楚的技术亮点。"
    )
    add_heading(document, "9.2 难题二：大表格会把文本规模瞬间放大", level=2)
    add_body_paragraph(
        document,
        "Excel 文件在业务场景中很常见，但直接把所有行列拼成文本非常危险。"
        "不仅 token 数暴涨，而且相邻行容易互相干扰，模型会把上一行的类别拼到下一行的数值里。"
        "系统最终采用“保留表格二维结构 + 同时构造行段落 + 限量采样”的组合方案：前端继续可以看整表预览，后端抽取则只处理可控规模的数据。"
    )
    add_heading(document, "9.3 难题三：同一列存在多个候选值，容易归并错行", level=2)
    add_body_paragraph(
        document,
        "比如某一列是“数值”，某一行可能在多个段落和多个单元格里都能找到数字。"
        "如果没有主键约束，系统会把不同对象的值合并到同一行。项目最后的解决办法是："
        "先以首列建立行，再围绕该行已知主键搜集同列候选；评分时同时检查同一 table_id/source_row、同一 paragraph_id、主键值是否相同，以及上下文窗口是否出现当前行键值。"
        "只有当候选在这些维度上足够一致，才允许填入。"
    )
    add_heading(document, "9.4 难题四：模型可能返回非 JSON 或低质量 JSON", level=2)
    add_body_paragraph(
        document,
        "在真实工程里，模型输出格式不稳定是常见问题。llm_client 因此先做 fenced JSON 提取，再做 safe_load_json。"
        "一旦解析失败，系统不会报废，而是返回空字典并走规则回退。"
        "这种设计虽然看起来保守，但它保证了服务不会因为一次模型输出异常而中断整条任务链路。"
    )
    add_heading(document, "9.5 难题五：长任务、页面刷新和中断恢复", level=2)
    add_body_paragraph(
        document,
        "文档处理不是瞬时操作，尤其是批量文件或大表格。"
        "如果用户刷新页面或线程池中的任务出现暂时停滞，系统仍需可恢复。"
        "因此前端用 localStorage 保存最近任务摘要，后端则用 RECOVERY_STALE_SECONDS=5 的策略识别长期未更新任务并自动重新入队。"
        "这让系统更接近真实生产工具，而不是一次性的脚本演示。"
    )
    add_heading(document, "9.6 最终方案的意义", level=2)
    add_bullets(
        document,
        [
            "解决了长文本直接输入导致的失配、串值和整段复制问题。",
            "把不可解释的模型黑盒结果，改造成可追踪、可回放、可展示候选过程的工程系统。",
            "形成了“模型负责理解，规则负责约束，数据库负责记忆，前端负责解释”的稳定分工。"
        ],
    )


def add_conclusion(document: Document) -> None:
    document.add_heading("10 结论与后续优化方向", level=1)
    add_body_paragraph(
        document,
        "从当前仓库实现来看，DocFusion-AI 已经完成了一个较完整的文档理解闭环："
        "用户可以定义结果列、上传多个文件、查看后台进度、获得结构化结果、追踪来源段落或单元格，并导出最终表格。"
        "更重要的是，项目没有停留在“把模型接进来”这一步，而是围绕模型不稳定、长文本失配、多记录归并和结果解释性，做了多层工程化设计。"
    )
    add_heading(document, "10.1 当前系统已经具备的能力", level=2)
    add_bullets(
        document,
        [
            "前端具备完整的列配置、上传、预览、轮询、展示、导出、溯源交互。",
            "后端具备缓存复用、后台流水线、任务恢复、结果裁剪和字段级存储。",
            "抽取引擎具备字段语义初始化、分批识别、逐列归并、候选解释和模型回退。",
            "匹配引擎具备标准字段映射、向量缓存和不适用时自动跳过。",
        ],
    )
    add_heading(document, "10.2 后续可持续优化的方向", level=2)
    add_bullets(
        document,
        [
            "把字段字典、别名和说明配置做成后台可维护资源，而不是只放在代码资产文件中。",
            "补充更多文档类型支持，例如 PDF 与扫描件 OCR 管线，并继续复用当前的统一中间态。",
            "把 __field_options__ 和 __decision_trace__ 进一步可视化，让用户能看到“系统为什么选了这个值而不是另一个值”。",
            "在批量任务场景引入更正式的队列与任务监控，但前提仍然是保留当前这套可解释的流水线分层。"
        ],
    )
    add_callout(
        document,
        "一句话总结项目技术路线",
        [
            "DocFusion-AI 不是把整份文档一次性交给大模型，而是先解析结构、再清洗内容、按批识别候选、逐列完成主键归并、最后做标准字段匹配与结果溯源，"
            "从而解决长文本难以稳定匹配的问题，并把结果做成真正可落地使用的结构化数据。",
        ],
        fill="EAF6F5",
        border=THEME["teal"],
    )


def build_document() -> Document:
    document = Document()
    configure_styles(document)

    add_hero(document)
    add_image_with_caption(document, SCREENSHOTS["manual"], "图 0 现有新手指导手册界面风格（本 Word 版排版参考）", width=6.1)
    add_page_break(document)

    add_contents(document)
    add_page_break(document)

    add_project_overview(document)
    add_architecture(document)
    add_frontend(document)
    add_backend(document)
    add_parser_and_cleaning(document)
    add_extraction(document)
    add_matching(document)
    add_llm_usage(document)
    add_challenges(document)
    add_conclusion(document)

    section = document.sections[-1]
    section.start_type = WD_SECTION_START.CONTINUOUS
    return document


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
