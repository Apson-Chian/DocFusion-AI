import argparse
import json
import re
import uuid
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    from python_calamine import CalamineWorkbook

    CALAMINE_AVAILABLE = True
except ImportError:
    CALAMINE_AVAILABLE = False


def column_letter(index: int) -> str:
    index += 1
    letters = []
    while index > 0:
        index, remainder = divmod(index - 1, 26)
        letters.append(chr(65 + remainder))
    return "".join(reversed(letters))


class DocumentParser:
    def __init__(self):
        self.supported_extensions = {".docx", ".xlsx", ".md", ".txt"}

    @staticmethod
    def _clean_text(text: Any) -> str:
        if text is None:
            return ""
        text = str(text)
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    def _clean_rows(self, rows: list[list[Any]]) -> list[list[str]]:
        cleaned_rows: list[list[str]] = []
        for row in rows:
            cleaned = [self._clean_text(cell) for cell in row]
            if any(cleaned):
                cleaned_rows.append(cleaned)
        return cleaned_rows

    def _build_sheet_paragraphs(self, sheet_name: str, rows: list[list[str]]) -> list[str]:
        if not rows:
            return []

        header = []
        for row in rows:
            filled = [cell for cell in row if cell]
            if len(filled) >= 2:
                header = row
                break

        paragraphs = [
            f"工作表 {sheet_name}，共 {len(rows)} 行，{max((len(row) for row in rows), default=0)} 列。"
        ]
        if header:
            paragraphs.append("表头：" + " | ".join(cell for cell in header if cell)[:500])

        preview_pairs = []
        for row in rows[1:4]:
            pairs = []
            for index, value in enumerate(row):
                if not value:
                    continue
                key = header[index] if index < len(header) and header[index] else f"列{index + 1}"
                pairs.append(f"{key}={value}")
                if len(pairs) >= 6:
                    break
            if pairs:
                preview_pairs.append("；".join(pairs))
        paragraphs.extend(preview_pairs)
        return paragraphs

    def _build_table_view(
        self,
        table_id: str,
        rows: list[list[Any]],
        title: str | None = None,
        sheet_name: str | None = None,
    ) -> dict[str, Any] | None:
        cleaned_rows = self._clean_rows(rows)
        if not cleaned_rows:
            return None

        row_payload = []
        max_cols = max((len(row) for row in cleaned_rows), default=0)
        for row_index, row in enumerate(cleaned_rows):
            cells = []
            for col_index in range(max_cols):
                value = row[col_index] if col_index < len(row) else ""
                cells.append(
                    {
                        "row_index": row_index,
                        "col_index": col_index,
                        "locator": f"{column_letter(col_index)}{row_index + 1}",
                        "value": value,
                    }
                )
            row_payload.append({"row_index": row_index, "cells": cells})

        header = []
        for row in cleaned_rows:
            filled = [cell for cell in row if cell]
            if len(filled) >= 2:
                header = row
                break

        return {
            "table_id": table_id,
            "title": title or table_id,
            "sheet_name": sheet_name,
            "row_count": len(cleaned_rows),
            "column_count": max_cols,
            "header": header,
            "rows": row_payload,
        }

    def parse(self, file_path: Path, doc_id: str | None = None) -> dict[str, Any]:
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在：{file_path}")

        ext = file_path.suffix.lower()
        if ext not in self.supported_extensions:
            raise ValueError(f"不支持的文件格式：{ext}")

        if ext == ".docx":
            paragraphs, tables, table_views = self._parse_docx(file_path)
        elif ext == ".xlsx":
            paragraphs, tables, table_views = self._parse_xlsx(file_path)
        elif ext == ".md":
            paragraphs, tables, table_views = self._parse_markdown(file_path)
        else:
            paragraphs, tables, table_views = self._parse_txt(file_path)

        cleaned_paragraphs = [self._clean_text(item) for item in paragraphs if self._clean_text(item)]
        cleaned_tables = self._clean_rows_tables(tables)
        raw_text = "\n".join(cleaned_paragraphs)

        return {
            "doc_id": doc_id or str(uuid.uuid4()),
            "doc_type": ext.lstrip("."),
            "paragraphs": cleaned_paragraphs,
            "paragraph_items": [{"paragraph_index": idx, "text": text} for idx, text in enumerate(cleaned_paragraphs)],
            "tables": cleaned_tables,
            "table_views": table_views,
            "raw_text": raw_text,
        }

    def _clean_rows_tables(self, tables: list[list[list[Any]]]) -> list[list[list[str]]]:
        return [rows for rows in (self._clean_rows(table) for table in tables) if rows]

    def _parse_docx(self, file_path: Path) -> tuple[list[str], list[list[list[str]]], list[dict[str, Any]]]:
        if Document is None:
            raise ImportError("python-docx 未安装，无法解析 Word 文档")
        if not zipfile.is_zipfile(file_path):
            raise ValueError(f"文件 {file_path} 不是有效的 docx 压缩包")

        try:
            paragraphs, tables = self._parse_docx_with_python_docx(file_path)
        except KeyError:
            paragraphs, tables = self._parse_docx_manual(file_path)

        table_views = []
        for index, table in enumerate(tables, start=1):
            view = self._build_table_view(f"table_{index:03d}", table, title=f"Word 表格 {index}")
            if view:
                table_views.append(view)
        return paragraphs, tables, table_views

    def _parse_docx_with_python_docx(self, file_path: Path) -> tuple[list[str], list[list[list[str]]]]:
        doc = Document(str(file_path))
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if self._clean_text(paragraph.text)]
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append(
                    [
                        self._clean_text(cell.text)
                        or "\n".join(self._clean_text(paragraph.text) for paragraph in cell.paragraphs if self._clean_text(paragraph.text))
                        for cell in row.cells
                    ]
                )
            if any(any(cell for cell in row) for row in rows):
                tables.append(rows)
        return paragraphs, tables

    def _parse_docx_manual(self, file_path: Path) -> tuple[list[str], list[list[list[str]]]]:
        paragraphs: list[str] = []
        tables: list[list[list[str]]] = []
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        with zipfile.ZipFile(file_path, "r") as archive:
            with archive.open("word/document.xml") as stream:
                tree = ET.parse(stream)
                root = tree.getroot()

        for paragraph in root.findall(".//w:p", namespaces=ns):
            text = "".join(node.text for node in paragraph.findall(".//w:t", namespaces=ns) if node.text)
            if self._clean_text(text):
                paragraphs.append(text)

        for table in root.findall(".//w:tbl", namespaces=ns):
            rows = []
            for row in table.findall(".//w:tr", namespaces=ns):
                values = []
                for cell in row.findall(".//w:tc", namespaces=ns):
                    parts = []
                    for paragraph in cell.findall(".//w:p", namespaces=ns):
                        text = "".join(node.text for node in paragraph.findall(".//w:t", namespaces=ns) if node.text)
                        if self._clean_text(text):
                            parts.append(text)
                    values.append("\n".join(parts))
                if any(self._clean_text(value) for value in values):
                    rows.append(values)
            if rows:
                tables.append(rows)

        return paragraphs, tables

    def _parse_xlsx(self, file_path: Path) -> tuple[list[str], list[list[list[str]]], list[dict[str, Any]]]:
        if openpyxl is not None:
            try:
                workbook = openpyxl.load_workbook(file_path, data_only=True)
                if workbook.worksheets:
                    return self._parse_openpyxl_workbook(workbook)
            except Exception:
                pass

        if not CALAMINE_AVAILABLE:
            raise ImportError("无法解析该 Excel 文件，请安装 python-calamine")

        workbook = CalamineWorkbook.from_path(str(file_path))
        paragraphs: list[str] = []
        tables: list[list[list[str]]] = []
        table_views: list[dict[str, Any]] = []

        for index, sheet_name in enumerate(workbook.sheet_names, start=1):
            sheet = workbook.get_sheet_by_name(sheet_name)
            rows = [[self._clean_text(cell) for cell in row] for row in sheet.to_python()]
            rows = [row for row in rows if any(row)]
            if not rows:
                continue
            tables.append(rows)
            paragraphs.extend(self._build_sheet_paragraphs(sheet_name, rows))
            view = self._build_table_view(
                f"sheet_{index:03d}",
                rows,
                title=sheet_name,
                sheet_name=sheet_name,
            )
            if view:
                table_views.append(view)

        return paragraphs, tables, table_views

    def _parse_openpyxl_workbook(self, workbook) -> tuple[list[str], list[list[list[str]]], list[dict[str, Any]]]:
        paragraphs: list[str] = []
        tables: list[list[list[str]]] = []
        table_views: list[dict[str, Any]] = []

        for index, sheet in enumerate(workbook.worksheets, start=1):
            rows = []
            for row in sheet.iter_rows(values_only=True):
                values = [self._clean_text(cell) for cell in row]
                if any(values):
                    rows.append(values)
            if not rows:
                continue
            tables.append(rows)
            paragraphs.extend(self._build_sheet_paragraphs(sheet.title, rows))
            view = self._build_table_view(
                f"sheet_{index:03d}",
                rows,
                title=sheet.title,
                sheet_name=sheet.title,
            )
            if view:
                table_views.append(view)

        return paragraphs, tables, table_views

    def _parse_markdown(self, file_path: Path) -> tuple[list[str], list[list[list[str]]], list[dict[str, Any]]]:
        content = self._read_text_file(file_path)
        paragraphs: list[str] = []
        tables: list[list[list[str]]] = []
        table_views: list[dict[str, Any]] = []
        lines = content.splitlines()
        index = 0

        while index < len(lines):
            line = lines[index].rstrip("\n")
            if line.strip().startswith("|"):
                rows = []
                while index < len(lines) and lines[index].strip().startswith("|"):
                    rows.append([cell.strip() for cell in lines[index].strip("| \n").split("|")])
                    index += 1
                if rows:
                    tables.append(rows)
                    view = self._build_table_view(
                        f"table_{len(table_views) + 1:03d}",
                        rows,
                        title=f"Markdown 表格 {len(table_views) + 1}",
                    )
                    if view:
                        table_views.append(view)
                continue
            if line.strip().startswith("#"):
                paragraphs.append(line.lstrip("#").strip())
            elif line.strip():
                paragraphs.append(line.strip())
            index += 1

        return paragraphs, tables, table_views

    def _parse_txt(self, file_path: Path) -> tuple[list[str], list[list[list[str]]], list[dict[str, Any]]]:
        content = self._read_text_file(file_path)
        paragraphs = [line.strip() for line in content.splitlines() if line.strip()]
        return paragraphs, [], []

    def _read_text_file(self, file_path: Path) -> str:
        for encoding in ["utf-8", "gbk", "gb2312", "utf-16", "latin-1"]:
            try:
                return file_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        raise UnicodeDecodeError("unknown", b"", 0, 1, f"无法解码文件: {file_path}")


def batch_parse(input_dir: Path, output_dir: Path | None = None, parser: DocumentParser | None = None):
    input_dir = Path(input_dir).resolve()
    output_dir = Path(output_dir).resolve() if output_dir else input_dir / "output_json"
    output_dir.mkdir(parents=True, exist_ok=True)
    parser = parser or DocumentParser()

    for file_path in input_dir.rglob("*"):
        if not file_path.is_file() or file_path.suffix.lower() not in parser.supported_extensions:
            continue
        rel_path = file_path.relative_to(input_dir)
        dir_part = "_".join(rel_path.parts[:-1])
        stem = file_path.stem
        suffix = file_path.suffix.lower().lstrip(".")
        doc_id = f"{dir_part + '_' if dir_part else ''}{stem}_{suffix}"
        result = parser.parse(file_path, doc_id=doc_id)
        (output_dir / f"{doc_id}.json").write_text(
            json.dumps(result, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )


def main():
    cli = argparse.ArgumentParser(description="批量解析文档")
    cli.add_argument("input_dir", type=str)
    cli.add_argument("-o", "--output_dir", type=str, default=None)
    args = cli.parse_args()
    batch_parse(Path(args.input_dir), Path(args.output_dir) if args.output_dir else None)


if __name__ == "__main__":
    main()
